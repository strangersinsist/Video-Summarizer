import os
import shutil
import time

import numpy as np
import streamlit as st
import tempfile
import cv2
from PIL import Image
import io
from docx import Document
from docx.shared import Inches
from streamlit_drawable_canvas import st_canvas

from sideBar import sidebar_image_processing
from transcript import transcribe_audio
from transcript import transcribe_audio
from audioRecord import audio_record


class VideoNoteApp:
    def __init__(self):
        self.initialize_session_state()

    def initialize_session_state(self):
        if "snapshots" not in st.session_state:
            st.session_state.snapshots = []
        if "video_file" not in st.session_state:
            st.session_state.video_file = None
        if "edited_image" not in st.session_state:
            st.session_state.edited_image = None
        if "current_frame_pos" not in st.session_state:
            st.session_state.current_frame_pos = 0
        if "canvas_key" not in st.session_state:
            st.session_state.canvas_key = 0
        if "url" not in st.session_state:
            st.session_state.url = 0
        if "api_key" not in st.session_state:
            st.session_state.api_key = "sk-104b5033564947be8d1ff74cc154b43a"
        if "points" not in st.session_state:
            st.session_state.points = []
        if "labels" not in st.session_state:
            st.session_state.labels = []

    def add_snapshot(self, image_data, video_time, annotation):
        st.session_state.snapshots.append({"image": image_data, "time": video_time, "annotation": annotation})

    def delete_snapshot(self, index):
        del st.session_state.snapshots[index]

    def update_annotation(self, index, new_annotation):
        st.session_state.snapshots[index]["annotation"] = new_annotation

    def export_to_docx(self):
        doc = Document()
        doc.add_heading('Video Notes', 0)
        for snapshot in st.session_state.snapshots:
            doc.add_picture(io.BytesIO(snapshot["image"]), width=Inches(4))
            doc.add_paragraph(f"Time: {snapshot['time']} seconds")
            doc.add_paragraph(snapshot["annotation"])
            doc.add_paragraph("\n")
        return doc

    def load_video(self):
        st.markdown("## Upload Video File")
        uploaded_video = st.file_uploader("Select video file", type=["mp4", "mov", "avi"])

        if uploaded_video is not None:
            # Save the uploaded video file
            tfile = tempfile.NamedTemporaryFile(delete=False)
            tfile.write(uploaded_video.read())
            st.session_state.video_file = tfile.name

            # Get the video file name and extract the video ID
            video_filename = uploaded_video.name
            video_id = video_filename.split('.')[0]  # Get the filename without the extension
            youtube_url = f"https://www.youtube.com/watch?v={video_id}"
            st.session_state.url = youtube_url  # Save the complete YouTube URL

        # Display the uploaded video
        if st.session_state.video_file is not None:
            st.video(st.session_state.video_file)

    def take_notes(self):
        if st.session_state.video_file is not None:
            cap = cv2.VideoCapture(st.session_state.video_file)
            if cap.isOpened():
                st.success("Video loaded successfully")

                total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                fps = int(cap.get(cv2.CAP_PROP_FPS))
                duration = total_frames / fps

                seconds_slider = st.slider("Select video time (seconds)", 0, int(duration), 0)
                frame_pos = int(seconds_slider * fps)

                if seconds_slider != st.session_state.current_frame_pos / fps:
                    st.session_state.current_frame_pos = frame_pos
                    st.session_state.canvas_key += 1
                    st.session_state.edited_image = None

                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_pos)
                ret, frame = cap.read()

                if ret:
                    img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
                    st.session_state.edited_image = io.BytesIO()
                    img.save(st.session_state.edited_image, format="PNG")
                    st.session_state.edited_image.seek(0)

                    st.image(img, caption=f"Current Frame: {int(seconds_slider)} seconds", use_column_width=True)
                    st.markdown(f"**Video Time: {int(seconds_slider)} seconds**")

                    # 调用侧边栏图像处理功能
                    sidebar_image_processing(img)


                    drawing_mode = st.selectbox("Select drawing mode",
                                                ("freedraw", "line", "rect", "circle", "transform"))
                    stroke_width = st.slider("Brush width: ", 1, 25, 3)
                    color = st.color_picker("Pick a color: ", "#000000")

                    canvas_result = st_canvas(
                        fill_color="rgba(255, 165, 0, 0.3)",
                        stroke_width=stroke_width,
                        stroke_color=color,
                        background_image=Image.open(st.session_state.edited_image).convert(
                            "RGBA") if st.session_state.edited_image else None,
                        update_streamlit=True,
                        height=360,
                        drawing_mode=drawing_mode,
                        key=f"canvas_{st.session_state.canvas_key}",
                    )

                    if canvas_result.image_data is not None:
                        edited_image = Image.fromarray(canvas_result.image_data.astype("uint8"), "RGBA")
                        original_img = Image.open(st.session_state.edited_image).convert("RGBA")

                        if original_img.size != edited_image.size:
                            edited_image = edited_image.resize(original_img.size)

                        combined_image = Image.alpha_composite(original_img, edited_image)

                        img_byte_arr = io.BytesIO()
                        combined_image.save(img_byte_arr, format="PNG")
                        st.session_state.edited_image = img_byte_arr.getvalue()



            # 初始化 session state
            if 'recording' not in st.session_state:
                st.session_state.recording = False
            if 'annotation' not in st.session_state:
                st.session_state.annotation = ""
            if 'transcribed' not in st.session_state:
                st.session_state.transcribed = False

            # 创建文本输入框
            annotation_placeholder = st.empty()
            annotation = annotation_placeholder.text_input("Enter Annotation", value=st.session_state.annotation,
                                                           key="text_input")
            # 创建录音按钮
            start_button = st.button('Start Recording')
            # 如果按下开始按钮，则开始录音
            if start_button:
                st.session_state.recording = True
                st.session_state.transcribed = False  # 当开始新的录音时，重置 transcribed 标志

            # 如果正在录音，则开始录音过程
            if st.session_state.recording:
                file_path = 'video/audio.wav'
                audio_record(file_path)
                countdown_placeholder = st.empty()
                for i in range(10, 0, -1):
                    countdown_placeholder.markdown(f"Notice,The record will end after {i} seconds")
                    time.sleep(1)
                countdown_placeholder.empty()  # 清除倒计时文本
                st.session_state.recording = False  # 停止录音
                ##st.success("Audio saved successfully.")
            audio_path = "video/audio.wav"

            # 定义一个函数来处理转录结
            def handle_transcription():
                if not st.session_state.transcribed and not st.session_state.recording:  # 只有在没有录音且尚未转录的情况下才进行转录
                    recognized_text = transcribe_audio(audio_path)
                    st.session_state.transcribed = True

                    pcm_file_path1 = "speech-vad-demo/pcm/16k_1.pcm"
                    if os.path.exists(pcm_file_path1):
                        os.remove(pcm_file_path1)
                    output_pcm_dir = ".\\speech-vad-demo\\output_pcm\\"
                    if os.path.exists(output_pcm_dir) and os.path.isdir(output_pcm_dir):
                        shutil.rmtree(output_pcm_dir)
                        os.makedirs(output_pcm_dir)

                    return recognized_text
                return None

            # 调用处理转录结果的函数
            recognized_text = handle_transcription()
            if st.session_state.transcribed and recognized_text is not None:
                st.markdown(
                    f"Transcript result: **<span style='color: red;'>{recognized_text}</span>**，You can choose to copy it into the annotation.",
                    unsafe_allow_html=True)

            if st.button("Add annotation") and st.session_state.edited_image and annotation:
                self.add_snapshot(st.session_state.edited_image, int(seconds_slider), annotation)
                st.success(f"Annotation added: {int(seconds_slider)} seconds - {annotation}")

            self.display_notes()


    def display_notes(self):
        st.markdown("### Added Annotations")
        with st.container():
            for i, snapshot in enumerate(st.session_state.snapshots):
                with st.expander(f"Annotation {i + 1}: {snapshot['time']} seconds"):
                    st.image(snapshot["image"], caption=f"{snapshot['time']} seconds - {snapshot['annotation']}")
                    new_annotation = st.text_area(f"Edit Annotation {i + 1}", value=snapshot["annotation"],
                                                  key=f"edit_{i}")
                    if st.button("Update Annotation", key=f"update_{i}"):
                        self.update_annotation(i, new_annotation)
                        st.success("Annotation updated")
                    if st.button("Delete", key=f"delete_{i}"):
                        self.delete_snapshot(i)
                        st.experimental_rerun()


    def export_notes(self):
        st.markdown("## Export Notes")
        if st.button("Export as DOCX"):
            doc = self.export_to_docx()
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("Download DOCX File", data=buffer, file_name="video_notes.docx")
